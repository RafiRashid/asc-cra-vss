import streamlit as st
import pdfplumber
import pandas as pd
import docx

st.set_page_config(
    page_title="ASC - Données VSS des CRA",
    page_icon="🧊",
    layout="wide",
    initial_sidebar_state="expanded",
    menu_items={
        'Get Help': 'https://docs.streamlit.io/',
        #'Report a bug': "rafi.rashid-abdur@service-civique.gouv.fr",
        'About': "Cette app permet de passer de déposer des comptes-rendu annuel (CRA) et de récupérer les données sur les VSS (réponses aux questions)."
    }
)

question = "Existe-t-il un plan de prévention des VSS au sein de votre organisme ?"

# Fonctions pour extraire tout le texte d'un fichier PDF ou WORD
def extract_text_from_pdf(pdf_file):
    with pdfplumber.open(pdf_file) as pdf:
        full_text = ""
        for page in pdf.pages:
            full_text += page.extract_text()
    return full_text

def extract_text_from_pdf_table(pdf_file):
    with pdfplumber.open(pdf_file) as pdf:
        first_page = pdf.pages[0]
        text = first_page.extract_text()
        nom_orga = [line.split("Nom de l’organisme")[1].strip() for line in text.split('\n') if "Nom de l’organisme" in line]
        num_agrement = ["NA-" + line.split("NA-")[1].strip() for line in text.split('\n') if "NA-" in line]
        return nom_orga, num_agrement

def extract_text_from_docx(docx_file):
    doc = docx.Document(docx_file)
    full_text = "\n".join([para.text for para in doc.paragraphs])
    full_text = full_text.replace('\xa0',' ')
    return full_text

def extract_text_from_docx_table(docx_file):
    doc = docx.Document(docx_file)
    table1 = doc.tables[0]
    l = []
    nom_orga = []
    for rows in table1.rows:
        for cells in rows.cells:
            l.append(cells.text)
    num_agrement = [l[2]]
    nom_orga = [l[5]]

    return nom_orga, num_agrement

def extract_info_from_text(full_text, question):
    
    #num_agrement = ["NA-" + full_text.split("NA-", 1)[1].split("\n",1)[0]]
    answers = ["","","",""]

    if question in full_text:
        after_question = full_text.split(question, 1)[1].strip()
        answers[0] = after_question.split("?")[0].split("Vos tuteurs",1)[0].replace('\n',' ').strip()
        answers[1] = after_question.split("?")[2].split("Vos volontaires",1)[0].replace('\n',' ').strip()
        answers[2] = after_question.split("?")[4].split("En cas de signalement",1)[0].replace('\n',' ').strip()
        answers[3] = after_question.split("?")[6].split("Page",1)[0].replace('\n',' ').strip()
            
    #return num_agrement, answers
    return answers


# Interface utilisateur avec Streamlit
st.title("Extraction des données sur les VSS")
st.write("Chargez plusieurs CRA au format PDF ou DOCX.")

# Champ pour télécharger plusieurs fichiers
uploaded_files = st.file_uploader("Déposez vos fichiers PDF ou DOCX ici", type=["pdf", "docx"], accept_multiple_files=True)

# Vérification et extraction
if uploaded_files:
    results = []
    
    for uploaded_file in uploaded_files:
        if uploaded_file.type == "application/pdf":
            full_text = extract_text_from_pdf(uploaded_file)
            nom_orga, num_agrement = extract_text_from_pdf_table(uploaded_file)
        elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            full_text = extract_text_from_docx(uploaded_file)
            nom_orga, num_agrement = extract_text_from_docx_table(uploaded_file)
        else:
            continue
        
        answers = extract_info_from_text(full_text, question)
        results.append(nom_orga + num_agrement + answers)
    
    # Création du DataFrame
    df = pd.DataFrame(results, columns=["Nom de l'organisme", 
                                        "Numéro de l'agrément",
                                        "Existe-t-il un plan de prévention des VSS au sein de votre organisme ?", 
                                        "Vos tuteurs ont-ils été sensibilisés à ce thème ? Votre organisme serait-il intéressé parun module dédié dans le cadre de l’accompagnement des tuteurs ?", 
                                        "Vos volontaires ont-ils été sensibilisés à ce thème ? Si oui, comment ?", 
                                        "En cas de signalement de VSS, avez-vous mis en place des procédures internes ?Comment avez-vous géré la situation ?"])
    st.dataframe(df)
    
    # Ajouter un bouton de téléchargement en Excel
    excel_file = "extraction_resultats.xlsx"
    df.to_excel(excel_file, index=False)
    with open(excel_file, "rb") as f:
        st.download_button("Télécharger les résultats en Excel", f, file_name=excel_file)
else:
    st.write("Aucun fichier téléchargé.")